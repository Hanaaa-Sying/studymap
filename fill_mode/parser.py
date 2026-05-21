import re
import pdfplumber


def extract_text_from_pdf(pdf_path: str) -> str:
    with pdfplumber.open(pdf_path) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def extract_first_pages(pdf_path: str, n: int = 25) -> str:
    with pdfplumber.open(pdf_path) as pdf:
        pages = pdf.pages[:n]
        return "\n".join(page.extract_text() or "" for page in pages)


def detect_toc(pdf_path: str) -> bool:
    """Return True if the document has a table of contents."""
    with pdfplumber.open(pdf_path) as pdf:
        sample = "\n".join(page.extract_text() or "" for page in pdf.pages[:5])

    # keyword match
    toc_keywords = re.compile(
        r"(目\s*录|contents|table\s+of\s+contents)", re.IGNORECASE
    )
    if toc_keywords.search(sample):
        return True

    # pattern: lines ending with a page number (e.g. "Chapter 1 .... 3" or "第一章…… 5")
    toc_line = re.compile(r".{4,}[\s.·…]{2,}\d{1,4}\s*$", re.MULTILINE)
    matches = toc_line.findall(sample)
    return len(matches) >= 4


def detect_toc_docx(docx_path: str) -> bool:
    from docx import Document
    doc = Document(docx_path)
    sample = "\n".join(p.text for p in doc.paragraphs[:60])
    toc_keywords = re.compile(
        r"(目\s*录|contents|table\s+of\s+contents)", re.IGNORECASE
    )
    if toc_keywords.search(sample):
        return True
    toc_line = re.compile(r".{4,}[\s.·…]{2,}\d{1,4}\s*$", re.MULTILINE)
    return len(toc_line.findall(sample)) >= 4


def extract_text_from_docx(docx_path: str) -> str:
    from docx import Document
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_first_paras_docx(docx_path: str, n: int = 200) -> str:
    from docx import Document
    doc = Document(docx_path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paras[:n])


def extract_text_from_outline(outline_path: str) -> str:
    with open(outline_path, "r", encoding="utf-8") as f:
        return f.read()


def chunk_text(text: str, chunk_size: int = 6000, overlap: int = 500) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
    return chunks
